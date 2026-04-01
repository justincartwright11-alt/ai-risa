# --- Helper: TBA name and bout detection ---
def is_tba_name(name):
    if not name:
        return True
    n = str(name).strip().lower()
    return n in {
        "tba",
        "tbd",
        "to be announced",
        "to be determined",
        "unknown",
        "opponent tba",
    }

def is_true_tba_bout(fighter_a_name, fighter_b_name):
    return is_tba_name(fighter_a_name) or is_tba_name(fighter_b_name)
# --- Helper: Premium fighter profile summary ---
def summarize_fighter_profile(profile, prediction=None, matchup_data=None):
    if not profile or profile.get('name', '').upper() == 'TBA':
        return "TBA opponent. Full tactical summary pending confirmation."
    parts = []
    name = profile.get('name', '')
    style = profile.get('style', '')
    stance = profile.get('stance', '')
    titles = profile.get('titles', [])
    weight = profile.get('weight_class', '')
    record = profile.get('record', '')
    if titles:
        parts.append(f"{name} holds {', '.join(titles)}.")
    if style:
        parts.append(f"A {style.lower()} stylist" + (f" ({stance})" if stance else "") + ".")
    elif stance:
        parts.append(f"Fights from a {stance} stance.")
    if weight:
        parts.append(f"Competes at {weight}.")
    if record:
        parts.append(f"Record: {record}.")
    # Add more cues as available
    if not parts:
        parts.append(f"{name}: No detailed profile data available.")
    return " ".join(parts)

# --- Helper: Resolve real bout identity ---
def resolve_bout_identity(matchup_data, event_fight_name, prediction, fighter_a_profile, fighter_b_profile):
    # Priority: matchup_data > event fight entry > prediction > profile file name
    fighter_a_name = None
    fighter_b_name = None
    bout_title = None
    # 1. matchup_data
    if matchup_data and 'fighters' in matchup_data and len(matchup_data['fighters']) == 2:
        fighter_a_name, fighter_b_name = matchup_data['fighters']
    # 2. event fight entry
    elif event_fight_name and 'vs.' in event_fight_name:
        left, right = [x.strip() for x in event_fight_name.split('vs.')]
        fighter_a_name, fighter_b_name = left, right
    # 3. prediction
    elif prediction and 'fighter_a_name' in prediction and 'fighter_b_name' in prediction:
        fighter_a_name = prediction['fighter_a_name']
        fighter_b_name = prediction['fighter_b_name']
    # 4. profile file names
    if not fighter_a_name and fighter_a_profile:
        fighter_a_name = fighter_a_profile.get('name')
    if not fighter_b_name and fighter_b_profile:
        fighter_b_name = fighter_b_profile.get('name')
    # Compose bout title
    if fighter_a_name and fighter_b_name:
        bout_title = f"{fighter_a_name} vs. {fighter_b_name}"
    else:
        bout_title = event_fight_name or "Unknown Bout"
    return fighter_a_name, fighter_b_name, bout_title

# --- Helper: Build fighter profile summary ---
def build_fighter_profile_summary(profile):
    if not profile:
        return ""
    parts = []
    if 'titles' in profile and profile['titles']:
        parts.append(f"Titles: {', '.join(profile['titles'])}")
    if 'weight_class' in profile:
        parts.append(f"Weight: {profile['weight_class']}")
    if 'style' in profile:
        parts.append(f"Style: {profile['style']}")
    if 'stance' in profile:
        parts.append(f"Stance: {profile['stance']}")
    if 'record' in profile:
        parts.append(f"Record: {profile['record']}")
    return "; ".join(parts)

# --- Helper: Build matchup contrast summary ---
def build_matchup_contrast(profile_a, profile_b, prediction, matchup_data):
    # Only for real matchups (not TBA)
    if not profile_a or not profile_b or profile_a.get('name','').upper() == 'TBA' or profile_b.get('name','').upper() == 'TBA':
        return {
            "style_contrast": "TBA",
            "pace_contrast": "TBA",
            "range_contrast": "TBA",
            "attrition_risk": "TBA",
            "volatility_level": "TBA",
            "likely_scoring_pattern": "TBA",
            "finish_pressure_side": "TBA",
            "matchup_contrast_summary": "TBA opponent. Full contrast pending."
        }
    a_style = profile_a.get('style', 'unknown style')
    b_style = profile_b.get('style', 'unknown style')
    a_titles = profile_a.get('titles', [])
    b_titles = profile_b.get('titles', [])
    a_weight = profile_a.get('weight_class', '')
    b_weight = profile_b.get('weight_class', '')
    # Style contrast
    style_contrast = f"{a_style} vs {b_style}"
    # Pace contrast
    pace_contrast = "Measured vs. aggressive" if 'pressure' in a_style.lower() or 'aggressor' in a_style.lower() or 'pressure' in b_style.lower() or 'aggressor' in b_style.lower() else "Technical vs. technical"
    # Range contrast
    range_contrast = "Long-range vs. inside" if 'boxer' in a_style.lower() or 'boxer' in b_style.lower() else "Mid-range battle"
    # Attrition risk
    attrition_risk = "High if pace surges" if 'pressure' in a_style.lower() or 'pressure' in b_style.lower() else "Moderate"
    # Volatility
    volatility_level = "Elevated" if 'aggressor' in a_style.lower() or 'aggressor' in b_style.lower() else "Controlled"
    # Scoring pattern
    likely_scoring_pattern = "Rounds likely to be captured by cleaner work and ring control."
    # Finish pressure
    finish_pressure_side = profile_a.get('name','') if 'stoppage' in prediction.get('method_tendency','').lower() else "Decision tendency"
    # Combined summary
    summary = f"{style_contrast} | {pace_contrast} | {range_contrast} | Attrition: {attrition_risk} | Volatility: {volatility_level} | {likely_scoring_pattern}"
    return {
        "style_contrast": style_contrast,
        "pace_contrast": pace_contrast,
        "range_contrast": range_contrast,
        "attrition_risk": attrition_risk,
        "volatility_level": volatility_level,
        "likely_scoring_pattern": likely_scoring_pattern,
        "finish_pressure_side": finish_pressure_side,
        "matchup_contrast_summary": summary
    }

# --- Canonical fight intelligence object builder ---
def build_fight_report(prediction, fighter_a_profile, fighter_b_profile, matchup_data, event_summary, event_fight_name):
    # 1. Resolve real bout identity
    fighter_a_name, fighter_b_name, bout_title = resolve_bout_identity(matchup_data, event_fight_name, prediction, fighter_a_profile, fighter_b_profile)
    # 2. Only build premium section for non-TBA bouts (use strict TBA detection)
    if is_true_tba_bout(fighter_a_name, fighter_b_name):
        return None
    # 3. Build real fighter profile summaries
    fighter_a_profile_summary = summarize_fighter_profile(fighter_a_profile, prediction, matchup_data)
    fighter_b_profile_summary = summarize_fighter_profile(fighter_b_profile, prediction, matchup_data)
    # 4. Build matchup contrast
    matchup_contrast = build_matchup_contrast(fighter_a_profile, fighter_b_profile, prediction, matchup_data)
    # 5. Hydrate premium tactical fields (to be replaced with bout-specific synthesis)
    premium_fields, premium_sources = hydrate_premium_tactical_fields(
        prediction=prediction,
        fighter_a_profile=fighter_a_profile,
        fighter_b_profile=fighter_b_profile,
        matchup_data=matchup_data,
        event_summary=event_summary,
    )
    # 6. Round-flow projection split
    round_flow = premium_fields.get("round_flow_projection") or ""
    round_flow_obj = {"early": "", "middle": "", "late": ""}
    if round_flow and ";" in round_flow:
        parts = [p.strip() for p in round_flow.split(";")]
        if len(parts) == 3:
            round_flow_obj = {"early": parts[0], "middle": parts[1], "late": parts[2]}
    # 7. Corner instructions split
    corner_instructions = premium_fields.get("corner_instructions") or ""
    corner_obj = {"fighter_a": "", "fighter_b": ""}
    if corner_instructions and "|" in corner_instructions:
        ca, cb = [x.strip() for x in corner_instructions.split("|", 1)]
        corner_obj = {"fighter_a": ca, "fighter_b": cb}
    # 8. Visual metrics placeholder (to be filled by future logic)
    visual_metrics = {}
    # 9. Scorecard scenarios placeholder
    scorecard_scenarios = prediction.get("scorecard_scenarios") or ""
    # 10. Danger zones as list
    dz = premium_fields.get("danger_zones")
    if dz and isinstance(dz, str) and ";" in dz:
        dz_list = [z.strip() for z in dz.split(";") if z.strip()]
    else:
        dz_list = [dz] if dz else []
    return {
        "bout_title": bout_title,
        "fighter_a_name": fighter_a_name,
        "fighter_b_name": fighter_b_name,
        "fighter_a_profile_summary": fighter_a_profile_summary,
        "fighter_b_profile_summary": fighter_b_profile_summary,
        "matchup_contrast_summary": matchup_contrast["matchup_contrast_summary"],
        "predicted_winner": prediction.get("predicted_winner", ""),
        "method_tendency": prediction.get("method_tendency", ""),
        "confidence": prediction.get("confidence"),
        "decision_structure": premium_fields.get("decision_structure", ""),
        "energy_use": premium_fields.get("energy_use", ""),
        "fatigue_failure_points": premium_fields.get("fatigue_failure_points", ""),
        "mental_condition": premium_fields.get("mental_condition", ""),
        "collapse_triggers": premium_fields.get("collapse_triggers", ""),
        "main_tactical_edge": premium_fields.get("main_tactical_edge", ""),
        "round_flow_projection": round_flow_obj,
        "corner_instructions": corner_obj,
        "danger_zones": dz_list,
        "scorecard_scenarios": scorecard_scenarios,
        "visual_metrics": visual_metrics,
    }

# --- Premium markdown renderer from fight_report ---
def render_full_prim_fight_section(fight_report):
    lines = []
    lines.append(f"### {fight_report['bout_title']}")
    lines.append("")
    lines.append(f"**Predicted winner:** {fight_report['predicted_winner']}")
    lines.append(f"**Method tendency:** {fight_report['method_tendency']}")
    if fight_report.get('confidence') is not None:
        lines.append(f"**Confidence:** {fight_report['confidence']}")
    lines.append("")
    lines.append(f"**Fighter A profile:** {fight_report['fighter_a_profile_summary']}")
    lines.append(f"**Fighter B profile:** {fight_report['fighter_b_profile_summary']}")
    if fight_report['matchup_contrast_summary']:
        lines.append("")
        lines.append(f"**Matchup contrast:** {fight_report['matchup_contrast_summary']}")
    lines.append("")
    if fight_report['decision_structure']:
        lines.append(f"**Decision structure:** {fight_report['decision_structure']}")
    if fight_report['energy_use']:
        lines.append(f"**Energy use:** {fight_report['energy_use']}")
    if fight_report['fatigue_failure_points']:
        lines.append(f"**Fatigue failure points:** {fight_report['fatigue_failure_points']}")
    if fight_report['mental_condition']:
        lines.append(f"**Mental condition:** {fight_report['mental_condition']}")
    if fight_report['collapse_triggers']:
        lines.append(f"**Collapse triggers:** {fight_report['collapse_triggers']}")
    if fight_report['main_tactical_edge']:
        lines.append(f"**Main tactical edge:** {fight_report['main_tactical_edge']}")
    # Round-flow
    rf = fight_report['round_flow_projection']
    if any(rf.values()):
        lines.append("**Round-flow projection:**")
        if rf.get("early"): lines.append(f"- Early: {rf['early']}")
        if rf.get("middle"): lines.append(f"- Middle: {rf['middle']}")
        if rf.get("late"): lines.append(f"- Late: {rf['late']}")
    # Corner instructions
    ci = fight_report['corner_instructions']
    if any(ci.values()):
        lines.append("**Corner instructions:**")
        if ci.get("fighter_a"): lines.append(f"- {fight_report['fighter_a_name']}: {ci['fighter_a']}")
        if ci.get("fighter_b"): lines.append(f"- {fight_report['fighter_b_name']}: {ci['fighter_b']}")
    # Danger zones
    if fight_report['danger_zones']:
        lines.append("**Danger zones:**")
        for dz in fight_report['danger_zones']:
            lines.append(f"- {dz}")
    # Scorecard scenarios
    if fight_report['scorecard_scenarios']:
        lines.append(f"**Scorecard scenarios:** {fight_report['scorecard_scenarios']}")
    # Visual metrics (chart/table-ready, not rendered unless present)
    if fight_report['visual_metrics']:
        lines.append("**Visual metrics:** (see charts/tables)")
    lines.append("")
    return "\n".join(lines)
# Canonical section builder for Full Prim fight
def build_full_prim_fight_section(fight_context, matchup_data, fighter_a_profile, fighter_b_profile):
    lines = []
    fighter_a = matchup_data.get("fighter_a_name") or fight_context.get("fighter_a_name") or "Fighter A"
    fighter_b = matchup_data.get("fighter_b_name") or fight_context.get("fighter_b_name") or "Fighter B"
    lines.append(f"### {fighter_a} vs. {fighter_b}")
    lines.append("")
    lines.append(f"Predicted winner: {fight_context.get('predicted_winner', 'Unknown')}")
    lines.append(f"Method tendency: {fight_context.get('method_tendency', 'Unknown')}")
    if fight_context.get('confidence') is not None:
        lines.append(f"Confidence: {fight_context['confidence']}")
    lines.append("")
    ordered_fields = [
        ("decision_structure", "Decision structure"),
        ("energy_use", "Energy use"),
        ("fatigue_failure_points", "Fatigue failure points"),
        ("mental_condition", "Mental condition"),
        ("collapse_triggers", "Collapse triggers"),
        ("main_tactical_edge", "Main tactical edge"),
        ("round_flow_projection", "Round-flow projection"),
        ("corner_instructions", "Corner instructions"),
        ("danger_zones", "Danger zones"),
    ]
    for key, label in ordered_fields:
        value = (fight_context.get(key) or "").strip()
        if value:
            lines.append(f"{label}: {value}")
    lines.append("")
    return "\n".join(lines)
# Restore main() at top level
def main():
    parser = argparse.ArgumentParser()
    parser.add_argument('--mode', required=True, choices=['fighter', 'matchup', 'event'])
    parser.add_argument('--fighter', help='Fighter name for single dossier')
    parser.add_argument('--fighter-a', help='Fighter A for matchup')
    parser.add_argument('--fighter-b', help='Fighter B for matchup')
    parser.add_argument('--event', help='Event slug for event mode (e.g., boxing_wollongong_2026_04_05)')
    parser.add_argument('--report', help='Report type (Full Prim Report, Fighter/Trainer, Promoter, Broadcast, Premium Fan)')
    parser.add_argument('--format', help='Export format (DOCX or PDF)')
    args = parser.parse_args()

    valid_report_types = ["Full Prim Report", "Fighter/Trainer", "Promoter", "Broadcast", "Premium Fan"]
    report_slug_map = {
        "Full Prim Report": "full_prim_report",
        "Fighter/Trainer": "fighter_trainer",
        "Promoter": "promoter",
        "Broadcast": "broadcast",
        "Premium Fan": "premium_fan"
    }
    if not args.report or args.report not in valid_report_types:
        print("ERROR: --report argument is required and must be one of: Full Prim Report, Fighter/Trainer, Promoter, Broadcast, Premium Fan.")
        sys.exit(1)
    if not args.format or args.format.lower() not in ["docx", "pdf"]:
        print("ERROR: --format argument is required and must be DOCX or PDF.")
        sys.exit(1)

    # Mode selection logic (fighter/matchup/event)
    if args.mode == 'fighter':
        if not args.fighter:
            print('Missing --fighter argument.')
            sys.exit(1)
        passed = run_fighter_mode(args.fighter, args.report)
    elif args.mode == 'matchup':
        if not args.fighter_a or not args.fighter_b:
            print('Missing --fighter-a or --fighter-b argument.')
            sys.exit(1)
        passed = run_matchup_mode(args.fighter_a, args.fighter_b, args.report)
    elif args.mode == 'event':
        if not args.event:
            print('Missing --event argument.')
            sys.exit(1)
        passed = run_event_mode(args.event, args.report)
    else:
        print('Invalid mode.')
        sys.exit(1)

    if not passed:
        print('Workflow failed.')
        sys.exit(1)

    # Determine output paths for export logic
    if args.mode == 'event':
        if args.report == "Full Prim Report":
            base = args.event + "_full_prim_report"
        else:
            base = args.event
        docx_path = f"C:/ai_risa_data/reports/{base}.docx"
        pdf_path = f"C:/ai_risa_data/reports/{base}.pdf"
        # --- Full Prim Report: Real content population ---
        if args.report == "Full Prim Report":
            # 1. Read event summary JSON
            event_summary_path = f"C:/ai_risa_data/reports/{args.event}_summary.json"
            if not os.path.exists(event_summary_path):
                print(f"ERROR: Event summary not found: {event_summary_path}")
                sys.exit(1)
            with open(event_summary_path, 'r', encoding='utf-8') as f:
                event_summary = json.load(f)
            # 2. Build canonical report sections
            report_sections = []
            report_sections = []
            report_sections.append(f"# Full Prim Report: {event_summary.get('event','')}\n\n**Date:** {event_summary.get('date','')}  ")
            report_sections.append(f"**Venue:** {event_summary.get('venue','')}  ")
            report_sections.append(f"**Watch:** {event_summary.get('broadcast','')}\n\n---\n")
            report_sections.append("## Executive Summary\n\nThis AI-RISA master report provides a complete, event-wide breakdown for all stakeholders. It includes tactical, technical, and strategic analysis for every bout, plus event-wide trends and recommendations.\n\n---\n")
            tba_sections = []
            premium_sections = []
            for bout_key, bout in event_summary.get('results', {}).items():
                pred_path = bout.get('prediction_file')
                fight_name = os.path.splitext(os.path.basename(pred_path or bout_key))[0].replace('_prediction','').replace('_vs_',' vs. ').replace('_',' ').title()
                # Try to get fighter names from event JSON if possible
                import re
                parts = re.split(r"\s+vs\.\s+", fight_name.strip(), maxsplit=1, flags=re.IGNORECASE)
                left = parts[0].strip() if len(parts) > 0 else ""
                right = parts[1].strip() if len(parts) > 1 else ""
                if not pred_path or not os.path.exists(pred_path):
                    print(
                        "[FULL_PRIM_DEBUG]",
                        {
                            "fight_name": fight_name,
                            "pred_path": pred_path,
                            "left": left,
                            "right": right,
                            "fighter_a_name": None,
                            "fighter_b_name": None,
                            "bout_title": None,
                            "is_true_tba_bout": None,
                            "route": "tba_sections (no prediction)"
                        },
                    )
                    tba_sections.append(fight_name)
                    continue
                with open(pred_path, 'r', encoding='utf-8') as pf:
                    pred = json.load(pf)
                # Load fighter profiles and matchup data if available
                fighter_a_profile = None
                fighter_b_profile = None
                matchup_data = {}
                try:
                    fighter_a_profile_path = f"C:/ai_risa_data/fighters/{left.lower().replace(' ', '_')}.json"
                    if os.path.exists(fighter_a_profile_path):
                        with open(fighter_a_profile_path, 'r', encoding='utf-8') as f:
                            fighter_a_profile = json.load(f)
                    fighter_b_profile_path = f"C:/ai_risa_data/fighters/{right.lower().replace(' ', '_')}.json"
                    if os.path.exists(fighter_b_profile_path):
                        with open(fighter_b_profile_path, 'r', encoding='utf-8') as f:
                            fighter_b_profile = json.load(f)
                    matchup_path = f"C:/ai_risa_data/matchups/{left.lower().replace(' ', '_')}_vs_{right.lower().replace(' ', '_')}.json"
                    if os.path.exists(matchup_path):
                        with open(matchup_path, 'r', encoding='utf-8') as f:
                            matchup_data = json.load(f)
                except Exception:
                    pass
                # Resolve names first for TBA detection
                fighter_a_name, fighter_b_name, bout_title = resolve_bout_identity(matchup_data, fight_name, pred, fighter_a_profile, fighter_b_profile)
                tba_result = is_true_tba_bout(fighter_a_name, fighter_b_name)
                print(
                    "[FULL_PRIM_DEBUG]",
                    {
                        "fight_name": fight_name,
                        "pred_path": pred_path,
                        "left": left,
                        "right": right,
                        "fighter_a_name": fighter_a_name,
                        "fighter_b_name": fighter_b_name,
                        "bout_title": bout_title,
                        "is_true_tba_bout": tba_result,
                        "route": "tba_sections" if tba_result else "premium_sections"
                    },
                )
                if "lauren price" in (bout_title or "").lower():
                    print(
                        "[LAUREN_PRICE_DEBUG]",
                        {
                            "fighter_a_name": fighter_a_name,
                            "fighter_b_name": fighter_b_name,
                            "bout_title": bout_title,
                            "route": "tba_sections" if tba_result else "premium_sections",
                        },
                    )
                if tba_result:
                    tba_sections.append(bout_title)
                    continue
                fight_report = build_fight_report(
                    prediction=pred,
                    fighter_a_profile=fighter_a_profile,
                    fighter_b_profile=fighter_b_profile,
                    matchup_data=matchup_data,
                    event_summary=event_summary,
                    event_fight_name=fight_name,
                )
                if fight_report:
                    # --- Fighter Comparison Table ---
                    fighter_table = "| Fighter | Titles | Weight |\n|---|---|---|\n"
                    fighter_table += f"| {fighter_a_name} | {', '.join(fighter_a_profile.get('titles', [])) if fighter_a_profile else ''} | {fighter_a_profile.get('weight_class','') if fighter_a_profile else ''} |\n"
                    fighter_table += f"| {fighter_b_name} | {', '.join(fighter_b_profile.get('titles', [])) if fighter_b_profile else ''} | {fighter_b_profile.get('weight_class','') if fighter_b_profile else ''} |\n"
                    # --- Tactical Summary Table ---
                    tactical = fight_report.get('matchup_contrast', {})
                    tactical_table = "| Style Contrast | Pace | Range | Attrition | Volatility | Scoring | Finish Pressure |\n|---|---|---|---|---|---|---|\n"
                    tactical_table += f"| {tactical.get('style_contrast','')} | {tactical.get('pace_contrast','')} | {tactical.get('range_contrast','')} | {tactical.get('attrition_risk','')} | {tactical.get('volatility_level','')} | {tactical.get('likely_scoring_pattern','')} | {tactical.get('finish_pressure_side','')} |\n"
                    # --- Scorecard Scenarios Table ---
                    def build_scorecard_table(fight_report):
                        # Placeholder: Replace with real logic if available
                        return "| Scenario | Winner | Notes |\n|---|---|---|\n| Standard | {} | Most likely |\n".format(fight_report.get('predicted_winner',''))
                    # --- Radar, Round-Flow, Heatmap Metrics ---
                    def build_radar_metrics(fight_report):
                        # Placeholder: Replace with real logic if available
                        return "Radar: " + str(fight_report.get('radar_metrics','N/A'))
                    def build_round_flow_metrics(fight_report):
                        return "Round Flow: " + str(fight_report.get('round_flow','N/A'))
                    def build_heatmap_metrics(fight_report):
                        return "Heatmap: " + str(fight_report.get('danger_zones','N/A'))
                    # --- Section Markdown ---
                    section_md = f"""
### {bout_title}

**Predicted winner:** {fight_report.get('predicted_winner','')}
**Method tendency:** {fight_report.get('method_tendency','')}
**Confidence:** {fight_report.get('confidence','')}

{fighter_table}

{tactical_table}

**Decision structure:** {fight_report.get('decision_structure','')}
**Energy use:** {fight_report.get('energy_use','')}
**Fatigue failure points:** {fight_report.get('fatigue_failure_points','')}
**Mental condition:** {fight_report.get('mental_condition','')}
**Collapse triggers:** {fight_report.get('collapse_triggers','')}
**Main tactical edge:** {fight_report.get('main_tactical_edge','')}
**Danger zones:** {fight_report.get('danger_zones','')}

{build_scorecard_table(fight_report)}

{build_radar_metrics(fight_report)}

{build_round_flow_metrics(fight_report)}

{build_heatmap_metrics(fight_report)}

---\n"
                    premium_sections.append(section_md)
            for section_md in premium_sections:
                report_sections.append(section_md)
            print("[FINAL_SECTION_COUNTS]", {"premium_sections": len(premium_sections), "tba_sections": len(tba_sections)})
            # Incomplete/TBA matchups
            if tba_sections:
                report_sections.append("---\n\n## Incomplete/TBA Matchup Watchlist\n")
                for tba in tba_sections:
                    report_sections.append(f"- {tba}: Incomplete matchup — opponent not confirmed. No full tactical briefing available.")
            # Event-wide tactical conclusions
            report_sections.append("\n---\n\n## Event-Wide Conclusions\n- All scheduled bouts are projected to be competitive, with a tendency toward decisions.\n- Main event features a unified world title defense for Lauren Price; tactical edge and home advantage noted.\n- Undercard features rising prospects; monitor for late changes or TBA replacements.\n\n---\n")
            report_sections.append("## Recommendations\n- Promote the main event as a historic homecoming and multi-title defense.\n- Undercard can be marketed for prospect development and local fan engagement.\n- Monitor weigh-ins and late changes for tactical updates.\n\n---\n")
            report_sections.append("## Disclaimer\nThis report is for internal use by event stakeholders only. AI-RISA outputs are predictive and based on available data as of the report date. Not for public distribution or betting purposes.\n")
            full_report_md = "\n\n".join(report_sections)
            from pathlib import Path
            md_path = Path(f"C:/ai_risa_data/reports/{base}.md")
            print("[MD_WRITE_DEBUG]", {"md_path": str(md_path), "exists_before": md_path.exists()})
            if md_path.exists():
                md_path.unlink()
            md_path.write_text(full_report_md, encoding="utf-8")
            print("[MD_WRITE_PREVIEW]", full_report_md[:300])
            # DOCX output from markdown
            doc = Document()
            for para in full_report_md.split("\n\n"):
                if para.strip().startswith("#"):
                    level = para.count("#")
                    doc.add_heading(para.replace("#", "").strip(), level=level)
                elif para.strip():
                    doc.add_paragraph(para.strip())
            doc.save(docx_path)
        else:
            # Minimal content for other event reports
            doc = Document()
            doc.add_heading(f"AI-RISA Event Report", 0)
            doc.add_paragraph(f"Event: {args.event}")
            doc.add_paragraph(f"Report Type: {args.report}")
            doc.add_paragraph(f"Generated on: {__import__('datetime').datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
            doc.add_paragraph("\nThis is a real DOCX file generated by python-docx. Replace this with full report content as needed.")
            doc.save(docx_path)
    elif args.mode == 'fighter':
        slug = report_slug_map.get(args.report, args.report.lower().replace(' ', '_').replace('/', '_'))
        base = args.fighter.lower().replace(' ', '_') + "_dossier_" + slug
        docx_path = f"C:/ai_risa_data/reports/{base}.docx"
        pdf_path = f"C:/ai_risa_data/reports/{base}.pdf"
        # Minimal real DOCX for fighter mode
        doc = Document()
        doc.add_heading(f"AI-RISA Fighter Dossier", 0)
        doc.add_paragraph(f"Fighter: {args.fighter}")
        doc.add_paragraph(f"Report Type: {args.report}")
        doc.add_paragraph(f"Generated on: {__import__('datetime').datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        doc.add_paragraph("\nThis is a real DOCX file generated by python-docx. Replace this with full report content as needed.")
        doc.save(docx_path)
    elif args.mode == 'matchup':
        slug = report_slug_map.get(args.report, args.report.lower().replace(' ', '_').replace('/', '_'))
        base = args.fighter_a.lower().replace(' ', '_') + "_vs_" + args.fighter_b.lower().replace(' ', '_') + "_" + slug
        docx_path = f"C:/ai_risa_data/reports/{base}.docx"
        pdf_path = f"C:/ai_risa_data/reports/{base}.pdf"
        # Minimal real DOCX for matchup mode
        doc = Document()
        doc.add_heading(f"AI-RISA Matchup Report", 0)
        doc.add_paragraph(f"Matchup: {args.fighter_a} vs {args.fighter_b}")
        doc.add_paragraph(f"Report Type: {args.report}")
        doc.add_paragraph(f"Generated on: {__import__('datetime').datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        doc.add_paragraph("\nThis is a real DOCX file generated by python-docx. Replace this with full report content as needed.")
        doc.save(docx_path)
    else:
        print("Unknown mode for export.")
        sys.exit(1)

    export_format = args.format.strip().lower()
    if export_format in ["1", "docx", "docx.", "docx,", "docx "]:
        # Validate DOCX
        try:
            with open(docx_path, 'rb') as f:
                docx_header = f.read(2)
                f.seek(0, os.SEEK_END)
                docx_size = f.tell()
            if docx_size < 1024:
                print(f"DOCX validation failed: file too small ({docx_size} bytes). Path: {docx_path}")
                sys.exit(1)
            if docx_header != b'PK':
                print(f"DOCX validation failed: header is not PK. Path: {docx_path}")
                sys.exit(1)
            print(f"DOCX report generated and validated: {docx_path}")
            sys.exit(0)
        except Exception as e:
            print(f"DOCX validation error: {e}")
            sys.exit(1)
    elif export_format in ["2", "pdf", "pdf.", "pdf,", "pdf "]:
        # Validate DOCX first
        try:
            with open(docx_path, 'rb') as f:
                docx_header = f.read(2)
                f.seek(0, os.SEEK_END)
                docx_size = f.tell()
            if docx_size < 1024:
                print(f"DOCX validation failed: file too small ({docx_size} bytes). Path: {docx_path}")
                sys.exit(1)
            if docx_header != b'PK':
                print(f"DOCX validation failed: header is not PK. Path: {docx_path}")
                sys.exit(1)
        except Exception as e:
            print(f"DOCX validation error: {e}")
            sys.exit(1)
        # Check for PDF export script
        ps1_path = "C:/ai_risa_data/export_docx_to_pdf.ps1"
        if not os.path.exists(ps1_path):
            print(f"PDF export script missing: {ps1_path}\nPDF not generated. DOCX master remains at: {docx_path}")
            if os.path.exists(pdf_path):
                try: os.remove(pdf_path)
                except: pass
            sys.exit(1)
        # Call PowerShell script to export DOCX to PDF
        try:
            result = subprocess.run([
                "powershell", "-ExecutionPolicy", "Bypass", "-File",
                ps1_path,
                "-DocxPath", docx_path,
                "-PdfPath", pdf_path
            ], capture_output=True, text=True)
            if result.returncode != 0:
                print("PDF export failed:")
                print(result.stdout)
                print(result.stderr)
                print(f"DOCX master remains at: {docx_path}")
                if os.path.exists(pdf_path):
                    try: os.remove(pdf_path)
                    except: pass
                sys.exit(1)
            # Validate PDF header
            with open(pdf_path, 'rb') as pf:
                pdf_header = pf.read(5)
                pf.seek(0, os.SEEK_END)
                pdf_size = pf.tell()
            if pdf_size < 1024:
                print(f"PDF validation failed: file too small ({pdf_size} bytes). Path: {pdf_path}")
                if os.path.exists(pdf_path):
                    try: os.remove(pdf_path)
                    except: pass
                print(f"DOCX master remains at: {docx_path}")
                sys.exit(1)
            if pdf_header != b'%PDF-':
                print(f"PDF validation failed: header is not %PDF-. Path: {pdf_path}")
                if os.path.exists(pdf_path):
                    try: os.remove(pdf_path)
                    except: pass
                print(f"DOCX master remains at: {docx_path}")
                sys.exit(1)
            print(f"DOCX written to: {docx_path}")
            print(f"PDF written to:  {pdf_path}")
            print("PDF export and validation successful.")
        except Exception as e:
            print(f"PDF export error: {e}")
            print(f"DOCX master remains at: {docx_path}")
            if os.path.exists(pdf_path):
                try: os.remove(pdf_path)
                except: pass
            sys.exit(1)
    else:
        print("Unknown export format. Please choose DOCX or PDF.")
        sys.exit(1)

import argparse
import os
import sys
from docx import Document
import json
import subprocess

# Premium tactical fields for enrichment
PREMIUM_TACTICAL_FIELDS = [
    "decision_structure",
    "energy_use",
    "fatigue_failure_points",
    "mental_condition",
    "collapse_triggers",
    "main_tactical_edge",
    "round_flow_projection",
    "corner_instructions",
    "danger_zones",
]

def hydrate_premium_tactical_fields(prediction, fighter_a_profile, fighter_b_profile, matchup_data, event_summary):
    hydrated = {}
    sources = {}
    winner = prediction.get("predicted_winner", "")
    loser = fighter_b_profile.get("name", "Opponent") if fighter_a_profile and fighter_a_profile.get("name") == winner else fighter_a_profile.get("name", "Opponent") if fighter_b_profile else "Opponent"
    method = prediction.get("method_tendency", "")
    # Deterministic enrichment for each field
    for field in PREMIUM_TACTICAL_FIELDS:
        val = prediction.get(field)
        if val:
            hydrated[field] = val
            sources[field] = "direct"
            continue
        # Inference logic for each field
        if field == "decision_structure":
            if method.lower() == "decision":
                hydrated[field] = f"{winner} is projected to win through cleaner sequencing, superior positional discipline, and steadier round capture rather than volatility."
            else:
                hydrated[field] = f"{winner} is projected to win by imposing their preferred fight rhythm and capitalizing on key openings."
        elif field == "energy_use":
            hydrated[field] = f"{winner}'s likely path is measured early work, efficient exchanges, and controlled output preservation over the full distance."
        elif field == "fatigue_failure_points":
            hydrated[field] = f"The main risk appears if prolonged clinch disruption or forced attritional rhythm drags {winner} into a less efficient late tempo."
        elif field == "mental_condition":
            hydrated[field] = f"Structured, composed, and process-led; unlikely to chase chaos without reason."
        elif field == "collapse_triggers":
            hydrated[field] = f"Repeated physical disruption, broken rhythm, and sustained pressure at uncomfortable range are the main routes to structural erosion."
        elif field == "main_tactical_edge":
            hydrated[field] = f"Cleaner decision structure and round-by-round scoring control."
        elif field == "round_flow_projection":
            hydrated[field] = f"Early range establishment, mid-fight widening through cleaner work, late lead protection unless {loser} creates disruption."
        elif field == "corner_instructions":
            hydrated[field] = f"Establish lead-hand control early, deny rhythm breaks, do not overcommit once ahead, and reset immediately after clean scoring sequences."
        elif field == "danger_zones":
            hydrated[field] = f"Clinch disruption, urgency surges from a losing opponent, and any late tempo spike that forces reactive exchanges."
        else:
            hydrated[field] = ""
        sources[field] = "inferred"
    return hydrated, sources

def run_fighter_mode(fighter, report_type):
    fighter_file = f"C:/ai_risa_data/fighters/{fighter.lower().replace(' ', '_')}.json"
    if not os.path.exists(fighter_file):
        print(f"Fighter file missing: {fighter_file}. Please create or update the fighter file first.")
        return False
    # Load fighter data
    with open(fighter_file, 'r', encoding='utf-8') as f:
        data = json.load(f)
    # Generate dossier (minimal fields)
    dossier = {
        "name": data["name"],
        "nickname": data.get("nickname", ""),
        "weight_class": data.get("weight_class", ""),
        "style_summary": data.get("style", ""),
        "decision_structure": data.get("decision_structure", "Aggressive exchanges, risk-taking"),
        "energy_use": data.get("energy_use", "High early, fades late"),
        "fatigue_failure_points": data.get("fatigue_failure_points", "Late rounds, under pressure"),
        "mental_condition": data.get("mental_condition", "Streak-dependent"),
        "collapse_triggers": data.get("collapse_triggers", "Sustained body attack, pace"),
        "tactical_notes": data.get("notes", "")
    }
    safe_report_type = report_type.lower().replace(' ', '_').replace('/', '_')
    out_path = f"C:/ai_risa_data/reports/{fighter.lower().replace(' ', '_')}_dossier_{safe_report_type}.json"
    with open(out_path, 'w', encoding='utf-8') as f:
        json.dump(dossier, f, indent=2)
    print(f"Fighter dossier written: {out_path}")
    return True

def run_matchup_mode(fighter_a, fighter_b, report_type):
    # Ensure both fighter files exist
    for fighter in [fighter_a, fighter_b]:
        fighter_file = f"C:/ai_risa_data/fighters/{fighter.lower().replace(' ', '_')}.json"
        if not os.path.exists(fighter_file):
            print(f"Fighter file missing: {fighter_file}. Please create or update the fighter file first.")
            return False
    # Create or update matchup file
    matchup_slug = f"{fighter_a.lower().replace(' ', '_')}_vs_{fighter_b.lower().replace(' ', '_')}"
    matchup_file = f"C:/ai_risa_data/matchups/{matchup_slug}.json"
    if not os.path.exists(matchup_file):
        matchup = {
            "fighters": [fighter_a, fighter_b],
            "weight_class": "",
            "scheduled_rounds": 10,
            "title_fight": False,
            "event_date": "",
            "venue": "",
            "notes": "Auto-generated for head-to-head."
        }
        with open(matchup_file, 'w', encoding='utf-8') as f:
            json.dump(matchup, f, indent=2)
    # Use or create a minimal runner for head-to-head
    runner = "C:/ai_risa_data/run_head_to_head.py"
    if not os.path.exists(runner):
        with open(runner, 'w', encoding='utf-8') as f:
            f.write("""import sys, json\n\nf1, f2 = sys.argv[1], sys.argv[2]\nmatchup_file = f'C:/ai_risa_data/matchups/{f1.lower().replace(' ', '_')}_vs_{f2.lower().replace(' ', '_')}.json'\nwith open(matchup_file, 'r', encoding='utf-8') as f:\n    matchup = json.load(f)\n# Minimal prediction logic\npred = {\n    'predicted_winner': f1,\n    'method_tendency': 'Decision',\n    'confidence': 0.6,\n    'tactical_interaction': 'Standard AI-RISA logic',\n    'round_flow_projection': 'Likely to go full rounds'\n}\nout_path = f'C:/ai_risa_data/predictions/{f1.lower().replace(' ', '_')}_vs_{f2.lower().replace(' ', '_')}_prediction.json'\nwith open(out_path, 'w', encoding='utf-8') as f:\n    json.dump(pred, f, indent=2)\nprint(f'Prediction written: {out_path}')\n""")
    # Run the head-to-head runner
    subprocess.run([sys.executable, runner, fighter_a, fighter_b], check=True)
    print(f"Head-to-head prediction written for {fighter_a} vs {fighter_b}.")
    return True

def run_event_mode(event_slug, report_type):
    runner = f"C:/ai_risa_data/run_{event_slug}.py"
    if not os.path.exists(runner):
        print(f"Event runner missing: {runner}. Please create or update the event runner first.")
        return False
    subprocess.run([sys.executable, runner], check=True)
    print(f"Event runner executed: {runner}")
    return True

    parser = argparse.ArgumentParser()
    parser.add_argument('--mode', required=True, choices=['fighter', 'matchup', 'event'])
    parser.add_argument('--fighter', help='Fighter name for single dossier')
    parser.add_argument('--fighter-a', help='Fighter A for matchup')
    parser.add_argument('--fighter-b', help='Fighter B for matchup')
    parser.add_argument('--event', help='Event slug for event mode (e.g., boxing_wollongong_2026_04_05)')
    parser.add_argument('--report', help='Report type (Full Prim Report, Fighter/Trainer, Promoter, Broadcast, Premium Fan)')
    parser.add_argument('--format', help='Export format (DOCX or PDF)')
    args = parser.parse_args()

    valid_report_types = ["Full Prim Report", "Fighter/Trainer", "Promoter", "Broadcast", "Premium Fan"]
    report_slug_map = {
        "Full Prim Report": "full_prim_report",
        "Fighter/Trainer": "fighter_trainer",
        "Promoter": "promoter",
        "Broadcast": "broadcast",
        "Premium Fan": "premium_fan"
    }
    if not args.report or args.report not in valid_report_types:
        print("ERROR: --report argument is required and must be one of: Full Prim Report, Fighter/Trainer, Promoter, Broadcast, Premium Fan.")
        sys.exit(1)
    if not args.format or args.format.lower() not in ["docx", "pdf"]:
        print("ERROR: --format argument is required and must be DOCX or PDF.")
        sys.exit(1)

    if args.mode == 'fighter':
        if not args.fighter:
            print('Missing --fighter argument.')
            sys.exit(1)
        passed = run_fighter_mode(args.fighter, args.report)
    elif args.mode == 'matchup':
        if not args.fighter_a or not args.fighter_b:
            print('Missing --fighter-a or --fighter-b argument.')
            sys.exit(1)
        passed = run_matchup_mode(args.fighter_a, args.fighter_b, args.report)
    elif args.mode == 'event':
        if not args.event:
            print('Missing --event argument.')
            sys.exit(1)
        passed = run_event_mode(args.event, args.report)
    else:
        print('Invalid mode.')
        sys.exit(1)


    if not passed:
        print('Workflow failed.')
        sys.exit(1)

    # Determine output paths for export logic
    if args.mode == 'event':
        if args.report == "Full Prim Report":
            base = args.event + "_full_prim_report"
        else:
            base = args.event
        docx_path = f"C:/ai_risa_data/reports/{base}.docx"
        pdf_path = f"C:/ai_risa_data/reports/{base}.pdf"
        # --- Full Prim Report: Real content population ---
        if args.report == "Full Prim Report":
            # 1. Read event summary JSON
            event_summary_path = f"C:/ai_risa_data/reports/{args.event}_summary.json"
            if not os.path.exists(event_summary_path):
                print(f"ERROR: Event summary not found: {event_summary_path}")
                sys.exit(1)
            with open(event_summary_path, 'r', encoding='utf-8') as f:
                event_summary = json.load(f)
            # 2. Read all fight prediction JSONs and build premium sections
            fight_sections = []
            tba_sections = []
            missing_preds = []
            for bout_key, bout in event_summary.get('results', {}).items():
                pred_path = bout.get('prediction_file')
                fight_name = os.path.splitext(os.path.basename(pred_path or bout_key))[0].replace('_prediction','').replace('_vs_',' vs. ').replace('_',' ').title()
                # Try to get fighter names from event JSON if possible
                if 'vs.' in fight_name:
                    left, right = fight_name.split(' vs. ')
                else:
                    left, right = fight_name, ''
                if 'TBA' in right or 'TBA' in left or not pred_path or not os.path.exists(pred_path):
                    tba_sections.append(fight_name)
                    continue
                with open(pred_path, 'r', encoding='utf-8') as pf:
                    pred = json.load(pf)
                # Load fighter profiles and matchup data if available
                fighter_a_profile = None
                fighter_b_profile = None
                matchup_data = None
                try:
                    fighter_a_profile_path = f"C:/ai_risa_data/fighters/{left.lower().replace(' ', '_')}.json"
                    if os.path.exists(fighter_a_profile_path):
                        with open(fighter_a_profile_path, 'r', encoding='utf-8') as f:
                            fighter_a_profile = json.load(f)
                    fighter_b_profile_path = f"C:/ai_risa_data/fighters/{right.lower().replace(' ', '_')}.json"
                    if os.path.exists(fighter_b_profile_path):
                        with open(fighter_b_profile_path, 'r', encoding='utf-8') as f:
                            fighter_b_profile = json.load(f)
                    matchup_path = f"C:/ai_risa_data/matchups/{left.lower().replace(' ', '_')}_vs_{right.lower().replace(' ', '_')}.json"
                    if os.path.exists(matchup_path):
                        with open(matchup_path, 'r', encoding='utf-8') as f:
                            matchup_data = json.load(f)
                except Exception:
                    pass
                hydrated, sources = hydrate_premium_tactical_fields(pred, fighter_a_profile, fighter_b_profile, matchup_data, event_summary)
                section = []
                section.append(f"Predicted winner: {pred.get('predicted_winner','N/A')}")
                section.append(f"Method tendency: {pred.get('method_tendency','N/A')}")
                if pred.get('confidence') is not None:
                    section.append(f"Confidence: {pred['confidence']}")
                for k in PREMIUM_TACTICAL_FIELDS:
                    if hydrated.get(k):
                        section.append(f"{k.replace('_',' ').capitalize()}: {hydrated[k]}")
                fight_sections.append((fight_name, section))
            # 3. Build DOCX report body
            doc = Document()
            # Cover Page
            doc.add_heading("AI-RISA Full Prim Report", 0)
            doc.add_paragraph(f"Event: {event_summary.get('event','')}")
            doc.add_paragraph(f"Date: {event_summary.get('date','')}")
            doc.add_paragraph(f"Venue: {event_summary.get('venue','')}")
            doc.add_paragraph(f"Generated on: {__import__('datetime').datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
            doc.add_page_break()
            # Executive Summary
            doc.add_heading("Executive Summary", level=1)
            doc.add_paragraph(f"This report provides a premium tactical breakdown and AI-RISA predictions for all confirmed bouts on the event card. Incomplete or TBA matchups are listed separately.")
            doc.add_page_break()
            # Fight-by-fight premium briefings
            for fight_title, section_lines in fight_sections:
                doc.add_heading(fight_title, level=1)
                for line in section_lines:
                    doc.add_paragraph(line)
                doc.add_page_break()
            # Incomplete/TBA matchups
            if tba_sections:
                doc.add_heading("Incomplete/TBA Matchup Watchlist", level=1)
                for tba in tba_sections:
                    doc.add_paragraph(f"{tba}: Incomplete matchup — opponent not confirmed. No full tactical briefing available.")
                doc.add_page_break()
            # Event-wide tactical conclusions
            doc.add_heading("Event-wide Tactical Conclusions", level=1)
            doc.add_paragraph("AI-RISA event-wide tactical summary and key takeaways.")
            doc.add_page_break()
            # Disclaimer
            doc.add_heading("Disclaimer", level=2)
            doc.add_paragraph("This report is generated by AI-RISA for informational and tactical analysis purposes only. All predictions and tactical notes are for professional use.")
            doc.save(docx_path)
        else:
            # Minimal content for other event reports
            doc = Document()
            doc.add_heading(f"AI-RISA Event Report", 0)
            doc.add_paragraph(f"Event: {args.event}")
            doc.add_paragraph(f"Report Type: {args.report}")
            doc.add_paragraph(f"Generated on: {__import__('datetime').datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
            doc.add_paragraph("\nThis is a real DOCX file generated by python-docx. Replace this with full report content as needed.")
            doc.save(docx_path)
    elif args.mode == 'fighter':
        slug = report_slug_map.get(args.report, args.report.lower().replace(' ', '_').replace('/', '_'))
        base = args.fighter.lower().replace(' ', '_') + "_dossier_" + slug
        docx_path = f"C:/ai_risa_data/reports/{base}.docx"
        pdf_path = f"C:/ai_risa_data/reports/{base}.pdf"
        # Minimal real DOCX for fighter mode
        doc = Document()
        doc.add_heading(f"AI-RISA Fighter Dossier", 0)
        doc.add_paragraph(f"Fighter: {args.fighter}")
        doc.add_paragraph(f"Report Type: {args.report}")
        doc.add_paragraph(f"Generated on: {__import__('datetime').datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        doc.add_paragraph("\nThis is a real DOCX file generated by python-docx. Replace this with full report content as needed.")
        doc.save(docx_path)
    elif args.mode == 'matchup':
        slug = report_slug_map.get(args.report, args.report.lower().replace(' ', '_').replace('/', '_'))
        base = args.fighter_a.lower().replace(' ', '_') + "_vs_" + args.fighter_b.lower().replace(' ', '_') + "_" + slug
        docx_path = f"C:/ai_risa_data/reports/{base}.docx"
        pdf_path = f"C:/ai_risa_data/reports/{base}.pdf"
        # Minimal real DOCX for matchup mode
        doc = Document()
        doc.add_heading(f"AI-RISA Matchup Report", 0)
        doc.add_paragraph(f"Matchup: {args.fighter_a} vs {args.fighter_b}")
        doc.add_paragraph(f"Report Type: {args.report}")
        doc.add_paragraph(f"Generated on: {__import__('datetime').datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        doc.add_paragraph("\nThis is a real DOCX file generated by python-docx. Replace this with full report content as needed.")
        doc.save(docx_path)
    else:
        print("Unknown mode for export.")
        sys.exit(1)

    export_format = args.format.strip().lower()
    if export_format in ["1", "docx", "docx.", "docx,", "docx "]:
        # Validate DOCX
        try:
            with open(docx_path, 'rb') as f:
                docx_header = f.read(2)
                f.seek(0, os.SEEK_END)
                docx_size = f.tell()
            if docx_size < 1024:
                print(f"DOCX validation failed: file too small ({docx_size} bytes). Path: {docx_path}")
                sys.exit(1)
            if docx_header != b'PK':
                print(f"DOCX validation failed: header is not PK. Path: {docx_path}")
                sys.exit(1)
            print(f"DOCX report generated and validated: {docx_path}")
            sys.exit(0)
        except Exception as e:
            print(f"DOCX validation error: {e}")
            sys.exit(1)
    elif export_format in ["2", "pdf", "pdf.", "pdf,", "pdf "]:
        # Validate DOCX first
        try:
            with open(docx_path, 'rb') as f:
                docx_header = f.read(2)
                f.seek(0, os.SEEK_END)
                docx_size = f.tell()
            if docx_size < 1024:
                print(f"DOCX validation failed: file too small ({docx_size} bytes). Path: {docx_path}")
                sys.exit(1)
            if docx_header != b'PK':
                print(f"DOCX validation failed: header is not PK. Path: {docx_path}")
                sys.exit(1)
        except Exception as e:
            print(f"DOCX validation error: {e}")
            sys.exit(1)
        # Check for PDF export script
        ps1_path = "C:/ai_risa_data/export_docx_to_pdf.ps1"
        if not os.path.exists(ps1_path):
            print(f"PDF export script missing: {ps1_path}\nPDF not generated. DOCX master remains at: {docx_path}")
            if os.path.exists(pdf_path):
                try: os.remove(pdf_path)
                except: pass
            sys.exit(1)
        # Call PowerShell script to export DOCX to PDF
        try:
            result = subprocess.run([
                "powershell", "-ExecutionPolicy", "Bypass", "-File",
                ps1_path,
                "-DocxPath", docx_path,
                "-PdfPath", pdf_path
            ], capture_output=True, text=True)
            if result.returncode != 0:
                print("PDF export failed:")
                print(result.stdout)
                print(result.stderr)
                print(f"DOCX master remains at: {docx_path}")
                if os.path.exists(pdf_path):
                    try: os.remove(pdf_path)
                    except: pass
                sys.exit(1)
            # Validate PDF header
            with open(pdf_path, 'rb') as pf:
                pdf_header = pf.read(5)
                pf.seek(0, os.SEEK_END)
                pdf_size = pf.tell()
            if pdf_size < 1024:
                print(f"PDF validation failed: file too small ({pdf_size} bytes). Path: {pdf_path}")
                if os.path.exists(pdf_path):
                    try: os.remove(pdf_path)
                    except: pass
                print(f"DOCX master remains at: {docx_path}")
                sys.exit(1)
            if pdf_header != b'%PDF-':
                print(f"PDF validation failed: header is not %PDF-. Path: {pdf_path}")
                if os.path.exists(pdf_path):
                    try: os.remove(pdf_path)
                    except: pass
                print(f"DOCX master remains at: {docx_path}")
                sys.exit(1)
            print(f"DOCX written to: {docx_path}")
            print(f"PDF written to:  {pdf_path}")
            print("PDF export and validation successful.")
        except Exception as e:
            print(f"PDF export error: {e}")
            print(f"DOCX master remains at: {docx_path}")
            if os.path.exists(pdf_path):
                try: os.remove(pdf_path)
                except: pass
            sys.exit(1)
    else:
        print("Unknown export format. Please choose DOCX or PDF.")
        sys.exit(1)

if __name__ == "__main__":
    main()
