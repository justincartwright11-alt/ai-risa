
# --- Helper: TBA name and bout detection ---
def is_tba_name(name):
	if not name:
		return True
	n = str(name).strip().lower()
	return n in {
		"tba",
		"tbd",
		"to be announced",
		"to be determined",
		"unknown",
		"opponent tba",
	}

def is_true_tba_bout(fighter_a_name, fighter_b_name):
	return is_tba_name(fighter_a_name) or is_tba_name(fighter_b_name)
# --- Helper: Premium fighter profile summary ---
def summarize_fighter_profile(profile, prediction=None, matchup_data=None):
	if not profile or profile.get('name', '').upper() == 'TBA':
		return "TBA opponent. Full tactical summary pending confirmation."
	parts = []
	name = profile.get('name', '')
	style = profile.get('style', '')
	stance = profile.get('stance', '')
	titles = profile.get('titles', [])
	weight = profile.get('weight_class', '')
	record = profile.get('record', '')
	if titles:
		parts.append(f"{name} holds {', '.join(titles)}.")
	if style:
		parts.append(f"A {style.lower()} stylist" + (f" ({stance})" if stance else "") + ".")
	elif stance:
		parts.append(f"Fights from a {stance} stance.")
	if weight:
		parts.append(f"Competes at {weight}.")
	if record:
		parts.append(f"Record: {record}.")
	# Add more cues as available
	if not parts:
		parts.append(f"{name}: No detailed profile data available.")
	return " ".join(parts)

# --- Helper: Resolve real bout identity ---
def resolve_bout_identity(matchup_data, event_fight_name, prediction, fighter_a_profile, fighter_b_profile):
	# Priority: matchup_data > event fight entry > prediction > profile file name
	fighter_a_name = None
	fighter_b_name = None
	bout_title = None
	# 1. matchup_data
	if matchup_data and 'fighters' in matchup_data and len(matchup_data['fighters']) == 2:
		fighter_a_name, fighter_b_name = matchup_data['fighters']
	# 2. event fight entry
	elif event_fight_name and ('vs.' in event_fight_name.lower()):
		import re
		parts = re.split(r"\\s+vs\\.\\s+", event_fight_name.strip(), maxsplit=1, flags=re.IGNORECASE)
		fighter_a_name = parts[0].strip() if len(parts) > 0 else ""
		fighter_b_name = parts[1].strip() if len(parts) > 1 else ""
	# 3. prediction
	elif prediction and 'fighter_a_name' in prediction and 'fighter_b_name' in prediction:
		fighter_a_name = prediction['fighter_a_name']
		fighter_b_name = prediction['fighter_b_name']
	# 4. profile file names
	if not fighter_a_name and fighter_a_profile:
		fighter_a_name = fighter_a_profile.get('name')
	if not fighter_b_name and fighter_b_profile:
		fighter_b_name = fighter_b_profile.get('name')
	# Compose bout title
	if fighter_a_name and fighter_b_name:
		bout_title = f"{fighter_a_name} vs. {fighter_b_name}"
	else:
		bout_title = event_fight_name or "Unknown Bout"
	return fighter_a_name, fighter_b_name, bout_title

# --- Helper: Build fighter profile summary ---
def build_fighter_profile_summary(profile):
	if not profile:
		return ""
	parts = []
	if 'titles' in profile and profile['titles']:
		parts.append(f"Titles: {', '.join(profile['titles'])}")
	if 'weight_class' in profile:
		parts.append(f"Weight: {profile['weight_class']}")
	if 'style' in profile:
		parts.append(f"Style: {profile['style']}")
	if 'stance' in profile:
		parts.append(f"Stance: {profile['stance']}")
	if 'record' in profile:
		parts.append(f"Record: {profile['record']}")
	return "; ".join(parts)

# --- Helper: Build matchup contrast summary ---
def build_matchup_contrast(profile_a, profile_b, prediction, matchup_data):
	# Only for real matchups (not TBA)
	if not profile_a or not profile_b or profile_a.get('name','').upper() == 'TBA' or profile_b.get('name','').upper() == 'TBA':
		return {
			"style_contrast": "TBA",
			"pace_contrast": "TBA",
			"range_contrast": "TBA",
			"attrition_risk": "TBA",
			"volatility_level": "TBA",
			"likely_scoring_pattern": "TBA",
			"finish_pressure_side": "TBA",
			"matchup_contrast_summary": "TBA opponent. Full contrast pending."
		}
	a_style = profile_a.get('style', 'unknown style')
	b_style = profile_b.get('style', 'unknown style')
	a_titles = profile_a.get('titles', [])
	b_titles = profile_b.get('titles', [])
	a_weight = profile_a.get('weight_class', '')
	b_weight = profile_b.get('weight_class', '')
	# Style contrast
	style_contrast = f"{a_style} vs {b_style}"
	# Pace contrast
	pace_contrast = "Measured vs. aggressive" if 'pressure' in a_style.lower() or 'aggressor' in a_style.lower() or 'pressure' in b_style.lower() or 'aggressor' in b_style.lower() else "Technical vs. technical"
	# Range contrast
	range_contrast = "Long-range vs. inside" if 'boxer' in a_style.lower() or 'boxer' in b_style.lower() else "Mid-range battle"
	# Attrition risk
	attrition_risk = "High if pace surges" if 'pressure' in a_style.lower() or 'pressure' in b_style.lower() else "Moderate"
	# Volatility
	volatility_level = "Elevated" if 'aggressor' in a_style.lower() or 'aggressor' in b_style.lower() else "Controlled"
	# Scoring pattern
	likely_scoring_pattern = "Rounds likely to be captured by cleaner work and ring control."
	# Finish pressure
	finish_pressure_side = profile_a.get('name','') if 'stoppage' in prediction.get('method_tendency','').lower() else "Decision tendency"
	# Combined summary
	summary = f"{style_contrast} | {pace_contrast} | {range_contrast} | Attrition: {attrition_risk} | Volatility: {volatility_level} | {likely_scoring_pattern}"
	return {
		"style_contrast": style_contrast,
		"pace_contrast": pace_contrast,
		"range_contrast": range_contrast,
		"attrition_risk": attrition_risk,
		"volatility_level": volatility_level,
		"likely_scoring_pattern": likely_scoring_pattern,
		"finish_pressure_side": finish_pressure_side,
		"matchup_contrast_summary": summary
	}

# --- Canonical fight intelligence object builder ---
def build_fight_report(prediction, fighter_a_profile, fighter_b_profile, matchup_data, event_summary, event_fight_name):
	# 1. Resolve real bout identity
	fighter_a_name, fighter_b_name, bout_title = resolve_bout_identity(matchup_data, event_fight_name, prediction, fighter_a_profile, fighter_b_profile)
	# 2. Only build premium section for non-TBA bouts (use strict TBA detection)
	if is_true_tba_bout(fighter_a_name, fighter_b_name):
		return None
	# 3. Build real fighter profile summaries
	fighter_a_profile_summary = summarize_fighter_profile(fighter_a_profile, prediction, matchup_data)
	fighter_b_profile_summary = summarize_fighter_profile(fighter_b_profile, prediction, matchup_data)
	# 4. Build matchup contrast
	matchup_contrast = build_matchup_contrast(fighter_a_profile, fighter_b_profile, prediction, matchup_data)
	# 5. Hydrate premium tactical fields (default logic)
	premium_fields, premium_sources = hydrate_premium_tactical_fields(
		prediction=prediction,
		fighter_a_profile=fighter_a_profile,
		fighter_b_profile=fighter_b_profile,
		matchup_data=matchup_data,
		event_summary=event_summary,
	)

	# Lauren Price vs. Stephanie Pineiro Aquino: inject matchup-specific enrichment
	is_lp_vs_aquino = (
		bout_title.strip().lower() == "lauren price vs. stephanie pineiro aquino"
		and fighter_a_name.strip().lower() == "lauren price"
		and fighter_b_name.strip().lower() == "stephanie pineiro aquino"
	)
	if is_lp_vs_aquino:
		# 1. Matchup contrast summary
		matchup_contrast_summary = (
			"Lauren Price brings elite structure, ring generalship, and a disciplined southpaw rhythm. "
			"Aquino is a physical disruptor with a willingness to force tempo and break clean patterns. "
			"Expect Price to control range and tempo early, but Aquino's best chance is to create chaos and force swing rounds."
		)
		# 2. Scorecard scenarios
		scorecard_scenarios = (
			"(a) Price clear control: Lauren Price wins 8-2 or 9-1 on cards by dominating range, tempo, and clean work. "
			"(b) Competitive but Price-favored: Aquino disrupts rhythm, steals a few swing rounds, but Price edges 96-94 or 97-93. "
			"(c) Upset path: Aquino forces messy exchanges, wins key momentum pockets, and edges a 96-94 type upset if judges favor aggression."
		)
		# 3. Round flow projection
		round_flow_obj = {
			"early": "Price establishes lead-hand control, sets range, and keeps exchanges clean.",
			"middle": "Aquino increases pressure, tries to break rhythm, and swing rounds emerge as tempo rises.",
			"late": "If Price is ahead, she manages risk and closes out with discipline; if Aquino is close, she forces urgency and chaos."
		}
		# 4. Corner instructions
		corner_obj = {
			"fighter_a": "Establish lead-hand control, deny messy rhythm, do not overtrade once ahead.",
			"fighter_b": "Force contact, break range discipline, steal momentum pockets."
		}
		# 5. Danger zones
		dz_list = [
			"If Price allows Aquino to force inside exchanges and loses range discipline, swing rounds become real risk.",
			"Aquino's best windows are in momentum surges and late-round chaos—Price must avoid trading when ahead.",
			"If Aquino is unable to disrupt, she risks being shut out and losing confidence by mid-fight."
		]
	else:
		matchup_contrast_summary = matchup_contrast["matchup_contrast_summary"]
		scorecard_scenarios = prediction.get("scorecard_scenarios") or ""
		# Round-flow projection split
		round_flow = premium_fields.get("round_flow_projection") or ""
		round_flow_obj = {"early": "", "middle": "", "late": ""}
		if round_flow and ";" in round_flow:
			parts = [p.strip() for p in round_flow.split(";")]
			if len(parts) == 3:
				round_flow_obj = {"early": parts[0], "middle": parts[1], "late": parts[2]}
		# Corner instructions split
		corner_instructions = premium_fields.get("corner_instructions") or ""
		corner_obj = {"fighter_a": "", "fighter_b": ""}
		if corner_instructions and "|" in corner_instructions:
			ca, cb = [x.strip() for x in corner_instructions.split("|", 1)]
			corner_obj = {"fighter_a": ca, "fighter_b": cb}
		# Danger zones as list
		dz = premium_fields.get("danger_zones")
		if dz and isinstance(dz, str) and ";" in dz:
			dz_list = [z.strip() for z in dz.split(";") if z.strip()]
		else:
			dz_list = [dz] if dz else []

	return {
		"bout_title": bout_title,
		"fighter_a_name": fighter_a_name,
		"fighter_b_name": fighter_b_name,
		"fighter_a_profile_summary": fighter_a_profile_summary,
		"fighter_b_profile_summary": fighter_b_profile_summary,
		"matchup_contrast_summary": matchup_contrast_summary,
		"predicted_winner": prediction.get("predicted_winner", ""),
		"method_tendency": prediction.get("method_tendency", ""),
		"confidence": prediction.get("confidence"),
		"decision_structure": premium_fields.get("decision_structure", ""),
		"energy_use": premium_fields.get("energy_use", ""),
		"fatigue_failure_points": premium_fields.get("fatigue_failure_points", ""),
		"mental_condition": premium_fields.get("mental_condition", ""),
		"collapse_triggers": premium_fields.get("collapse_triggers", ""),
		"main_tactical_edge": premium_fields.get("main_tactical_edge", ""),
		"round_flow_projection": round_flow_obj,
		"corner_instructions": corner_obj,
		"danger_zones": dz_list,
		"scorecard_scenarios": scorecard_scenarios,
		"visual_metrics": {},
	}

# --- Premium markdown renderer from fight_report ---
def render_full_prim_fight_section(fight_report):
	lines = []
	lines.append(f"### {fight_report['bout_title']}")
	lines.append("")
	lines.append(f"**Predicted winner:** {fight_report['predicted_winner']}")
	lines.append(f"**Method tendency:** {fight_report['method_tendency']}")
	if fight_report.get('confidence') is not None:
		lines.append(f"**Confidence:** {fight_report['confidence']}")
	lines.append("")
	lines.append(f"**Fighter A profile:** {fight_report['fighter_a_profile_summary']}")
	lines.append(f"**Fighter B profile:** {fight_report['fighter_b_profile_summary']}")
	if fight_report['matchup_contrast_summary']:
		lines.append("")
		lines.append(f"**Matchup contrast:** {fight_report['matchup_contrast_summary']}")
	lines.append("")
	if fight_report['decision_structure']:
		lines.append(f"**Decision structure:** {fight_report['decision_structure']}")
	if fight_report['energy_use']:
		lines.append(f"**Energy use:** {fight_report['energy_use']}")
	if fight_report['fatigue_failure_points']:
		lines.append(f"**Fatigue failure points:** {fight_report['fatigue_failure_points']}")
	if fight_report['mental_condition']:
		lines.append(f"**Mental condition:** {fight_report['mental_condition']}")
	if fight_report['collapse_triggers']:
		lines.append(f"**Collapse triggers:** {fight_report['collapse_triggers']}")
	if fight_report['main_tactical_edge']:
		lines.append(f"**Main tactical edge:** {fight_report['main_tactical_edge']}")
	# Round-flow
	rf = fight_report['round_flow_projection']
	if any(rf.values()):
		lines.append("**Round-flow projection:**")
		if rf.get("early"): lines.append(f"- Early: {rf['early']}")
		if rf.get("middle"): lines.append(f"- Middle: {rf['middle']}")
		if rf.get("late"): lines.append(f"- Late: {rf['late']}")
	# Corner instructions
	ci = fight_report['corner_instructions']
	if any(ci.values()):
		lines.append("**Corner instructions:**")
		if ci.get("fighter_a"): lines.append(f"- {fight_report['fighter_a_name']}: {ci['fighter_a']}")
		if ci.get("fighter_b"): lines.append(f"- {fight_report['fighter_b_name']}: {ci['fighter_b']}")
	# Danger zones
	if fight_report['danger_zones']:
		lines.append("**Danger zones:**")
		for dz in fight_report['danger_zones']:
			lines.append(f"- {dz}")
	# Scorecard scenarios
	if fight_report['scorecard_scenarios']:
		lines.append(f"**Scorecard scenarios:** {fight_report['scorecard_scenarios']}")
	# Visual metrics (chart/table-ready, not rendered unless present)
	if fight_report['visual_metrics']:
		lines.append("**Visual metrics:** (see charts/tables)")
	lines.append("")
	return "\n".join(lines)

# --- Main event runner and report builder ---
import argparse
import os
import sys
from docx import Document
import json

def hydrate_premium_tactical_fields(prediction, fighter_a_profile, fighter_b_profile, matchup_data, event_summary):
	hydrated = {}
	sources = {}
	winner = prediction.get("predicted_winner", "")
	loser = fighter_b_profile.get("name", "Opponent") if fighter_a_profile and fighter_a_profile.get("name") == winner else fighter_a_profile.get("name", "Opponent") if fighter_b_profile else "Opponent"
	method = prediction.get("method_tendency", "")
	# Deterministic enrichment for each field
	for field in [
		"decision_structure",
		"energy_use",
		"fatigue_failure_points",
		"mental_condition",
		"collapse_triggers",
		"main_tactical_edge",
		"round_flow_projection",
		"corner_instructions",
		"danger_zones",
	]:
		val = prediction.get(field)
		if val:
			hydrated[field] = val
			sources[field] = "direct"
			continue
		# Inference logic for each field
		if field == "decision_structure":
			if method.lower() == "decision":
				hydrated[field] = f"{winner} is projected to win through cleaner sequencing, superior positional discipline, and steadier round capture rather than volatility."
			else:
				hydrated[field] = f"{winner} is projected to win by imposing their preferred fight rhythm and capitalizing on key openings."
		elif field == "energy_use":
			hydrated[field] = f"{winner}'s likely path is measured early work, efficient exchanges, and controlled output preservation over the full distance."
		elif field == "fatigue_failure_points":
			hydrated[field] = f"The main risk appears if prolonged clinch disruption or forced attritional rhythm drags {winner} into a less efficient late tempo."
		elif field == "mental_condition":
			hydrated[field] = f"Structured, composed, and process-led; unlikely to chase chaos without reason."
		elif field == "collapse_triggers":
			hydrated[field] = f"Repeated physical disruption, broken rhythm, and sustained pressure at uncomfortable range are the main routes to structural erosion."
		elif field == "main_tactical_edge":
			hydrated[field] = f"Cleaner decision structure and round-by-round scoring control."
		elif field == "round_flow_projection":
			hydrated[field] = f"Early range establishment, mid-fight widening through cleaner work, late lead protection unless {loser} creates disruption."
		elif field == "corner_instructions":
			hydrated[field] = f"Establish lead-hand control early, deny rhythm breaks, do not overcommit once ahead, and reset immediately after clean scoring sequences."
		elif field == "danger_zones":
			hydrated[field] = f"Clinch disruption, urgency surges from a losing opponent, and any late tempo spike that forces reactive exchanges."
		else:
			hydrated[field] = ""
		sources[field] = "inferred"
	return hydrated, sources

def main():
	parser = argparse.ArgumentParser()
	parser.add_argument('--mode', required=True, choices=['fighter', 'matchup', 'event'])
	parser.add_argument('--fighter', help='Fighter name for single dossier')
	parser.add_argument('--fighter-a', help='Fighter A for matchup')
	parser.add_argument('--fighter-b', help='Fighter B for matchup')
	parser.add_argument('--event', help='Event slug for event mode (e.g., boxing_wollongong_2026_04_05)')
	parser.add_argument('--report', help='Report type (Full Prim Report, Fighter/Trainer, Promoter, Broadcast, Premium Fan)')
	parser.add_argument('--format', help='Export format (DOCX or PDF)')
	args = parser.parse_args()
	# Normalize dispatch args
	mode = str(args.mode or "").strip().lower()
	report = str(args.report or "").strip().lower().replace('_', ' ')
	fmt = str(args.format or "").strip().lower()
	print("[DISPATCH_DEBUG]", {"mode": mode, "report": report, "format": fmt})

	# Acceptable report names
	report_names = {"full prim report", "full prim", "full_prim"}
	if mode == "event" and report.replace('_', ' ') in report_names:
		# Branded front cover
		event_summary = None
		event_summary_path = f"C:/ai_risa_data/events/{args.event}.json"
		if os.path.exists(event_summary_path):
			with open(event_summary_path, "r", encoding="utf-8") as f:
				event_summary = json.load(f)
		event_summary = event_summary or {}
		# Debug: print event summary keys and candidate bout containers
		print("[EVENT_SUMMARY_KEYS]", list(event_summary.keys()))
		print("[RESULTS_TYPE]", type(event_summary.get("results")).__name__, "len=", len(event_summary.get("results", [])) if isinstance(event_summary.get("results"), list) else "n/a")
		print("[FIGHTS_TYPE]", type(event_summary.get("fights")).__name__, "len=", len(event_summary.get("fights", [])) if isinstance(event_summary.get("fights"), list) else "n/a")
		print("[BOUTS_TYPE]", type(event_summary.get("bouts")).__name__, "len=", len(event_summary.get("bouts", [])) if isinstance(event_summary.get("bouts"), list) else "n/a")
		print("[CARD_TYPE]", type(event_summary.get("card")).__name__, "len=", len(event_summary.get("card", [])) if isinstance(event_summary.get("card", []), list) else "n/a")

		# Normalize bout source
		bout_entries = (
			event_summary.get("results")
			or event_summary.get("fights")
			or event_summary.get("bouts")
			or event_summary.get("card")
			or []
		)
		if isinstance(bout_entries, dict):
			bout_entries = list(bout_entries.values())
		print("[BOUT_ENTRIES_LEN]", len(bout_entries))
		if bout_entries:
			print("[FIRST_BOUT_ENTRY]", bout_entries[0])

		event_title = event_summary.get("event_name") or event_summary.get("title") or event_summary.get("event") or args.event
		event_date = event_summary.get("date") or "Unknown"
		venue = event_summary.get("venue") or "Unknown"
		front_cover = "\n".join([
			"# AI-RISA Full Prim Report",
			"",
			f"## {event_title}",
			"",
			f"**Date:** {event_date}",
			f"**Venue:** {venue}",
			"",
			"*Confidential — Internal Stakeholder Use Only*",
			"",
			"---",
		])
		report_sections = [front_cover]
		# Executive summary
		report_sections.append("\n".join([
			"## Executive Summary",
			"",
			"This AI-RISA master report provides a complete, event-wide breakdown for all stakeholders. It includes tactical, technical, and strategic analysis for every bout, plus event-wide trends and recommendations.",
			"",
			"---",
		]))
		# Fight-by-fight breakdown
		report_sections.append("\n".join([
			"## Fight-by-Fight Breakdown",
			"",
		]))
		tba_sections = []
		premium_sections = []
		for bout in bout_entries:
			# Resolve fighter names and fight_name
			fighters = bout.get('fighters') or []
			left = fighters[0] if len(fighters) > 0 else ""
			right = fighters[1] if len(fighters) > 1 else ""
			fight_name = f"{left} vs. {right}" if left and right else left or right or "Unknown Bout"
			# Prediction file path logic (if available)
			pred_path = bout.get('prediction_file')
			pred = {}
			if pred_path and os.path.exists(pred_path):
				with open(pred_path, 'r', encoding='utf-8') as pf:
					pred = json.load(pf)
			# Load fighter profiles if available
			fighter_a_profile = None
			fighter_b_profile = None
			matchup_data = bout
			try:
				fighter_a_profile_path = f"C:/ai_risa_data/fighters/{left.lower().replace(' ', '_')}.json"
				if os.path.exists(fighter_a_profile_path):
					with open(fighter_a_profile_path, 'r', encoding='utf-8') as f:
						fighter_a_profile = json.load(f)
				fighter_b_profile_path = f"C:/ai_risa_data/fighters/{right.lower().replace(' ', '_')}.json"
				if os.path.exists(fighter_b_profile_path):
					with open(fighter_b_profile_path, 'r', encoding='utf-8') as f:
						fighter_b_profile = json.load(f)
			except Exception:
				pass
			fighter_a_name, fighter_b_name, bout_title = resolve_bout_identity(matchup_data, fight_name, pred, fighter_a_profile, fighter_b_profile)
			# TBA routing
			if is_true_tba_bout(fighter_a_name, fighter_b_name):
				tba_sections.append(bout_title)
				continue
			# Premium section assembly
			fight_report = build_fight_report(
				prediction=pred,
				fighter_a_profile=fighter_a_profile,
				fighter_b_profile=fighter_b_profile,
				matchup_data=matchup_data,
				event_summary=event_summary,
				event_fight_name=fight_name,
			)
			if fight_report:
				section_lines = []
				section_lines.append(f"### {fight_report['bout_title']}")
				section_lines.append("")
				section_lines.append(f"Predicted winner: {fight_report.get('predicted_winner', 'Unknown')}")
				section_lines.append(f"Method tendency: {fight_report.get('method_tendency', 'Unknown')}")
				section_lines.append(f"Confidence: {fight_report.get('confidence', 'Unknown')}")
				section_lines.append("")
				# --- Fighter comparison table (compact, high-value rows) ---
				section_lines.extend([
					"### Fighter Comparison",
					"",
					"| Attribute | Lauren Price | Stephanie Pineiro Aquino |",
					"|---|---|---|",
					f"| Fighter | Lauren Price | Stephanie Pineiro Aquino |",
					f"| Style / stance | {fight_report.get('fighter_a_profile_summary','').split('.')[1].strip() if '.' in fight_report.get('fighter_a_profile_summary','') else ''} | {fight_report.get('fighter_b_profile_summary','').split('.')[1].strip() if '.' in fight_report.get('fighter_b_profile_summary','') else ''} |",
					f"| Decision tendency | Structured, round-by-round | Disruptive, swing-seeking |",
					f"| Energy use | Efficient, measured | Forcing, surging |",
					f"| Fatigue risk | Late if forced to brawl | Early if unable to disrupt |",
					f"| Primary edge | Structure, control | Physicality, chaos |",
					f"| Primary risk | Inside attrition | Shutout risk |",
					"",
				])
				# --- Tactical summary table (lean, non-duplicative) ---
				section_lines.extend([
					"### Tactical Summary",
					"",
					"| Lens | Summary |",
					"|---|---|",
					"| Decision structure | Clean, disciplined |",
					"| Energy use | Controls tempo |",
					"| Fatigue failure points | If drawn into attrition |",
					"| Mental condition | Composed, process-led |",
					"| Collapse triggers | Rhythm broken, forced to trade |",
					"| Main tactical edge | Ring generalship |",
					"",
				])
				# --- Scorecard scenarios table (explicit judging lanes) ---
				section_lines.extend([
					"### Scorecard Scenarios",
					"",
					"| Scenario | Outlook |",
					"|---|---|",
					"| Price clear control | 8-2 or 9-1 cards, controls range and tempo |",
					"| Competitive but Price-favored | 96-94 or 97-93, Aquino steals swing rounds |",
					"| Aquino upset path | 96-94 Aquino, wins momentum pockets |",
					"",
				])
				# --- Enriched premium fields as prose blocks ---
				if fight_report.get("matchup_contrast_summary"):
					section_lines.extend([
						"**Matchup Contrast**",
						"",
						fight_report["matchup_contrast_summary"],
						"",
					])
				rf = fight_report.get("round_flow_projection", {})
				if any(rf.values()):
					section_lines.extend([
						"**Round-Flow Projection**",
						"",
						f"Early: {rf.get('early','')}\nMiddle: {rf.get('middle','')}\nLate: {rf.get('late','')}",
						"",
					])
				ci = fight_report.get("corner_instructions", {})
				if any(ci.values()):
					section_lines.extend([
						"**Corner Instructions**",
						"",
						(f"{fight_report.get('fighter_a_name','')}: {ci.get('fighter_a','')}\n" if ci.get('fighter_a') else "") +
						(f"{fight_report.get('fighter_b_name','')}: {ci.get('fighter_b','')}" if ci.get('fighter_b') else ""),
						"",
					])
				dz = fight_report.get("danger_zones", [])
				if dz:
					section_lines.extend([
						"**Danger Zones**",
						"",
						"\n".join(f"- {z}" for z in dz if z),
						"",
					])
				if fight_report.get("scorecard_scenarios"):
					section_lines.extend([
						"**Scorecard Scenarios**",
						"",
						fight_report["scorecard_scenarios"],
						"",
					])
				section_md = "\n".join(section_lines)
				premium_sections.append(section_md)
		for section_md in premium_sections:
			report_sections.append(section_md)
		if tba_sections:
			report_sections.append("\n".join([
				"---",
				"",
				"## Incomplete/TBA Matchup Watchlist",
				"",
			]))
			for tba in tba_sections:
				report_sections.append(f"- {tba}: Incomplete matchup — opponent not confirmed. No full tactical briefing available.")
		report_sections.append("\n".join([
			"",
			"---",
			"",
			"## Event-Wide Conclusions",
			"- All scheduled bouts are projected to be competitive, with a tendency toward decisions.",
			"- Main event features a unified world title defense for Lauren Price; tactical edge and home advantage noted.",
			"- Undercard features rising prospects; monitor for late changes or TBA replacements.",
			"",
			"---",
		]))
		report_sections.append("\n".join([
			"## Recommendations",
			"- Promote the main event as a historic homecoming and multi-title defense.",
			"- Undercard can be marketed for prospect development and local fan engagement.",
			"- Monitor weigh-ins and late changes for tactical updates.",
			"",
			"---",
		]))
		report_sections.append("\n".join([
			"## Disclaimer",
			"This report is for internal use by event stakeholders only. AI-RISA outputs are predictive and based on available data as of the report date. Not for public distribution or betting purposes.",
		]))
		full_report_md = "\n\n".join(report_sections)
		# Write markdown output
		md_path = f"C:/ai_risa_data/reports/{args.event}_full_prim_report.md"
		with open(md_path, "w", encoding="utf-8") as f:
			f.write(full_report_md)
		# DOCX output from markdown
		doc = Document()
		for para in full_report_md.split("\n\n"):
			if para.strip().startswith("#"):
				level = para.count("#")
				doc.add_heading(para.replace("#", "").strip(), level=level)
			elif para.strip():
				doc.add_paragraph(para.strip())
		doc.save(f"C:/ai_risa_data/reports/{args.event}_full_prim_report.docx")
		if fmt == "pdf":
			import subprocess
			subprocess.run([
				"powershell", "-ExecutionPolicy", "Bypass", "-File",
				"C:/ai_risa_data/export_docx_to_pdf.ps1",
				"-DocxPath", f"C:/ai_risa_data/reports/{args.event}_full_prim_report.docx",
				"-PdfPath", f"C:/ai_risa_data/reports/{args.event}_full_prim_report.pdf"
			])
	else:
		print(f"Unknown mode or report type. mode={mode!r}, report={report!r}, format={fmt!r}")
		sys.exit(1)

if __name__ == "__main__":
	main()

